#!/usr/bin/env python3
"""Build the canonical general and OpenAI-targeted CV artifacts."""

from __future__ import annotations

import os
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
SOURCE_DIR = REPO_ROOT / "src" / "data" / "resume"
OUTPUT_DIR = REPO_ROOT / "public" / "cv"

ARTIFACTS = (
    ("cv.md", "Dessi_Georgieva_CV", "General evidence-led CV"),
    (
        "openai-codex-cv.md",
        "Dessi_Georgieva_OpenAI_Codex_CV",
        "OpenAI Codex application CV",
    ),
)

LINK_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
TOKEN_PATTERN = re.compile(r"(\*\*[^*]+\*\*|_[^_]+_|\[[^\]]+\]\([^)]+\))")


def set_cell_free_page_geometry(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.32)


def set_repeatable_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.25)
    normal.font.color.rgb = RGBColor(35, 38, 42)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.04

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = RGBColor(20, 22, 24)
    title.paragraph_format.space_after = Pt(1)

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1.font.size = Pt(15)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(46, 116, 181)
    heading_1.paragraph_format.space_before = Pt(10)
    heading_1.paragraph_format.space_after = Pt(3)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2.font.size = Pt(11.5)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(46, 116, 181)
    heading_2.paragraph_format.space_before = Pt(6)
    heading_2.paragraph_format.space_after = Pt(1)
    heading_2.paragraph_format.keep_with_next = True

    if "CV Headline" not in styles:
        headline = styles.add_style("CV Headline", WD_STYLE_TYPE.PARAGRAPH)
    else:
        headline = styles["CV Headline"]
    headline.font.name = "Calibri"
    headline.font.size = Pt(11.5)
    headline.font.bold = True
    headline.font.color.rgb = RGBColor(92, 48, 44)
    headline.paragraph_format.space_after = Pt(3)

    if "CV Contact" not in styles:
        contact = styles.add_style("CV Contact", WD_STYLE_TYPE.PARAGRAPH)
    else:
        contact = styles["CV Contact"]
    contact.font.name = "Calibri"
    contact.font.size = Pt(8.7)
    contact.font.color.rgb = RGBColor(85, 89, 94)
    contact.paragraph_format.space_after = Pt(5)

    if "CV Meta" not in styles:
        meta = styles.add_style("CV Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["CV Meta"]
    meta.font.name = "Calibri"
    meta.font.size = Pt(9.8)
    meta.font.color.rgb = RGBColor(45, 48, 52)
    meta.paragraph_format.space_after = Pt(3)

    if "CV Bullet" not in styles:
        bullet = styles.add_style("CV Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = styles["CV Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(9.65)
    bullet.font.color.rgb = RGBColor(35, 38, 42)
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.12)
    bullet.paragraph_format.space_after = Pt(1.5)
    bullet.paragraph_format.line_spacing = 1.01


def set_document_metadata(document: Document, label: str) -> None:
    properties = document.core_properties
    properties.title = f"Dessi Georgieva - {label}"
    properties.subject = "AI systems engineering resume"
    properties.keywords = "AI systems, agents, evaluation, Python, FastAPI, TypeScript"
    properties.author = ""
    properties.last_modified_by = ""

    settings = document.settings.element
    language = settings.find(qn("w:themeFontLang"))
    if language is None:
        language = OxmlElement("w:themeFontLang")
        settings.append(language)
    language.set(qn("w:val"), "en-GB")


def add_bottom_rule(paragraph) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "E36B5D")
    borders.append(bottom)


def add_hyperlink(
    paragraph, text: str, url: str, *, bold: bool = False, italic: bool = False
) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    if italic:
        run_properties.append(OxmlElement("w:i"))
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_markdown(paragraph, value: str, *, italic: bool = False) -> None:
    position = 0
    for token in TOKEN_PATTERN.finditer(value):
        if token.start() > position:
            run = paragraph.add_run(value[position : token.start()])
            run.italic = italic
        raw = token.group(0)
        link = LINK_PATTERN.fullmatch(raw)
        if link:
            add_hyperlink(paragraph, link.group(1), link.group(2), italic=italic)
        elif raw.startswith("**"):
            run = paragraph.add_run(raw[2:-2])
            run.bold = True
            run.italic = italic
        elif raw.startswith("_"):
            add_inline_markdown(paragraph, raw[1:-1], italic=True)
        position = token.end()
    if position < len(value):
        run = paragraph.add_run(value[position:])
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 124, 128)


def build_docx(source: Path, target: Path, label: str) -> None:
    document = Document()
    set_cell_free_page_geometry(document)
    set_repeatable_styles(document)
    set_document_metadata(document, label)
    add_page_number(document.sections[0].footer.paragraphs[0])

    lines = source.read_text(encoding="utf-8").splitlines()
    seen_title = False
    seen_headline = False
    seen_contact = False

    for raw_line in lines:
        line = raw_line.strip()
        if not line or line == "---":
            continue

        if line.startswith("# ") and not seen_title:
            paragraph = document.add_paragraph(style="Title")
            paragraph.add_run(line[2:])
            add_bottom_rule(paragraph)
            seen_title = True
            continue

        if line.startswith("## "):
            if line in {"## Experience", "## Professional Experience"}:
                document.add_page_break()
            document.add_heading(line[3:], level=1)
            continue

        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="CV Bullet")
            paragraph.add_run("• ")
            add_inline_markdown(paragraph, line[2:])
            continue

        if seen_title and not seen_headline and line.startswith("**") and line.endswith("**"):
            paragraph = document.add_paragraph(style="CV Headline")
            paragraph.add_run(line[2:-2])
            seen_headline = True
            continue

        if seen_headline and not seen_contact:
            paragraph = document.add_paragraph(style="CV Contact")
            add_inline_markdown(paragraph, line)
            seen_contact = True
            continue

        paragraph = document.add_paragraph(
            style="CV Meta" if line.startswith("**") else None
        )
        if line.startswith("_") and line.endswith("_"):
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_markdown(paragraph, line[1:-1], italic=True)
        else:
            add_inline_markdown(paragraph, line)

    document.save(target)


def convert_to_pdf(docx_paths: list[Path]) -> None:
    configured = os.environ.get("SOFFICE_BIN")
    soffice = configured or shutil.which("soffice")
    if not soffice:
        print("DOCX and Markdown generated; PDF conversion skipped because soffice is unavailable.")
        return
    for docx_path in docx_paths:
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(OUTPUT_DIR),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                text=True,
            )
        except subprocess.CalledProcessError:
            print(
                f"PDF conversion skipped for {docx_path.name}; "
                "use the document render workflow to emit the verified PDF."
            )


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    generated_docx: list[Path] = []
    for source_name, stem, label in ARTIFACTS:
        source = SOURCE_DIR / source_name
        markdown_target = OUTPUT_DIR / f"{stem}.md"
        docx_target = OUTPUT_DIR / f"{stem}.docx"
        shutil.copyfile(source, markdown_target)
        build_docx(source, docx_target, label)
        generated_docx.append(docx_target)
        print(f"Generated {docx_target.relative_to(REPO_ROOT)}")
        print(f"Synced {markdown_target.relative_to(REPO_ROOT)}")
    convert_to_pdf(generated_docx)


if __name__ == "__main__":
    main()
